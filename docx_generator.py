# docx_generator.py — DOCX 보고서 생성 모듈
# python-docx 기반 / 분석 결과 자동 해석 텍스트 포함

from __future__ import annotations

import io
from datetime import datetime, date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np
import pandas as pd
from scipy import stats

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# python-docx import (Streamlit Cloud 호환)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 상수
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

TITLE_COLOR = RGBColor(0x1F, 0x4E, 0x79) if _DOCX_AVAILABLE else None   # 진청
ACCENT_COLOR = RGBColor(0x2E, 0x75, 0xB6) if _DOCX_AVAILABLE else None   # 파랑
HEADER_BG = "1F4E79"    # 헤더 배경 (hex, #없이)
ROW_BG_EVEN = "DEEAF1"  # 짝수 행 배경

CHART_DPI = 100
CHART_WIDTH = 6.5   # inches
CHART_HEIGHT_INCH = 3.5

# 한국어 폰트 설정 (나눔고딕 또는 malgun)
_KR_FONTS = ["NanumGothic", "Malgun Gothic", "AppleGothic", "DejaVu Sans"]


def _get_kr_font() -> str:
    available = {f.name for f in fm.fontManager.ttflist}
    for f in _KR_FONTS:
        if f in available:
            return f
    return "DejaVu Sans"


KR_FONT = _get_kr_font()


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 통계 계산 헬퍼
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _has(df: pd.DataFrame, col: str) -> bool:
    return col in df.columns and df[col].notna().any()


def _compute_summary(df: pd.DataFrame) -> dict:
    """주요 지표 요약 계산."""
    result: dict = {}

    if _has(df, "temp_avg"):
        result["평균기온"] = (df["temp_avg"].mean(), "°C")
    if _has(df, "temp_max"):
        result["최고기온 평균"] = (df["temp_max"].mean(), "°C")
        result["역대 최고기온"] = (df["temp_max"].max(), "°C")
        result["폭염일수 (≥33°C)"] = ((df["temp_max"] >= 33).sum(), "일")
    if _has(df, "temp_min"):
        result["최저기온 평균"] = (df["temp_min"].mean(), "°C")
        result["역대 최저기온"] = (df["temp_min"].min(), "°C")
        result["열대야일수 (≥25°C)"] = ((df["temp_min"] >= 25).sum(), "일")
        result["한파일수 (≤-12°C)"] = ((df["temp_min"] <= -12).sum(), "일")
    if _has(df, "precipitation"):
        result["연평균 강수량"] = (
            df.groupby("year")["precipitation"].sum().mean() if "year" in df.columns
            else df["precipitation"].sum(),
            "mm/년"
        )
        result["총 강수량"] = (df["precipitation"].sum(), "mm")
    if _has(df, "humidity"):
        result["평균 습도"] = (df["humidity"].mean(), "%")
    if _has(df, "wind_speed"):
        result["평균 풍속"] = (df["wind_speed"].mean(), "m/s")

    return result


def _compute_monthly_stats(df: pd.DataFrame) -> pd.DataFrame:
    """월별 평균 통계."""
    cols = [c for c in ["temp_avg", "temp_max", "temp_min", "precipitation",
                         "humidity", "wind_speed"] if _has(df, c)]
    if not cols or "month" not in df.columns:
        return pd.DataFrame()

    grp = ["month"]
    if "station_name" in df.columns and df["station_name"].nunique() == 1:
        pass  # 단일 관측소면 station_name 제외
    monthly = df.groupby(grp)[cols].mean().round(2).reset_index()

    rename = {
        "month": "월", "temp_avg": "평균기온(°C)", "temp_max": "최고기온(°C)",
        "temp_min": "최저기온(°C)", "precipitation": "강수량(mm)",
        "humidity": "습도(%)", "wind_speed": "풍속(m/s)"
    }
    return monthly.rename(columns=rename)


def _compute_annual_stats(df: pd.DataFrame) -> pd.DataFrame:
    """연도별 주요 통계."""
    if "year" not in df.columns:
        return pd.DataFrame()

    agg: dict = {}
    if _has(df, "temp_avg"):
        agg["평균기온(°C)"] = pd.NamedAgg("temp_avg", "mean")
    if _has(df, "precipitation"):
        agg["연강수량(mm)"] = pd.NamedAgg("precipitation", "sum")
    if _has(df, "temp_max"):
        agg["폭염일수"] = pd.NamedAgg("temp_max", lambda s: (s >= 33).sum())
    if _has(df, "temp_min"):
        agg["열대야일수"] = pd.NamedAgg("temp_min", lambda s: (s >= 25).sum())

    if not agg:
        return pd.DataFrame()

    grp = ["year"]
    if "station_name" in df.columns and df["station_name"].nunique() > 1:
        grp = ["year", "station_name"]

    return df.groupby(grp).agg(**agg).round(2).reset_index()


def _auto_interpret(df: pd.DataFrame) -> list[str]:
    """분석 결과 자동 해석 문장 생성."""
    lines: list[str] = []

    if "year" not in df.columns:
        return lines

    years = sorted(df["year"].dropna().unique())
    n_years = len(years)
    period = f"{int(years[0])}~{int(years[-1])}년 ({n_years}년간)"

    # 기온 추세
    if _has(df, "temp_avg"):
        annual_t = df.groupby("year")["temp_avg"].mean().dropna()
        if len(annual_t) >= 5:
            x = annual_t.index.values.astype(float)
            y = annual_t.values
            slope, intercept, r, p, _ = stats.linregress(x, y)
            sig = "통계적으로 유의한" if p < 0.05 else "유의하지 않은"
            direction = "상승" if slope > 0 else "하락"
            lines.append(
                f"[기온] {period} 평균기온은 10년당 {abs(slope * 10):.2f}°C "
                f"{direction} 추세로 {sig} 변화를 보임 (R²={r**2:.2f}, p={p:.3f})."
            )
            lines.append(
                f"  분석 기간 평균기온: {annual_t.mean():.1f}°C "
                f"(최저 {annual_t.min():.1f}°C [{int(annual_t.idxmin())}년], "
                f"최고 {annual_t.max():.1f}°C [{int(annual_t.idxmax())}년])"
            )

    # 폭염·열대야
    if _has(df, "temp_max") and "year" in df.columns:
        heatwave = df.groupby("year")["temp_max"].apply(lambda s: (s >= 33).sum())
        if len(heatwave) >= 5:
            x = heatwave.index.values.astype(float)
            slope, _, _, p, _ = stats.linregress(x, heatwave.values)
            sig = "증가" if (slope > 0 and p < 0.05) else "감소" if (slope < 0 and p < 0.05) else "변화 없음"
            lines.append(
                f"[폭염] 폭염일수(최고기온 ≥33°C)는 연평균 {heatwave.mean():.1f}일이며, "
                f"장기 추세는 {sig}."
            )

    # 강수 추세
    if _has(df, "precipitation"):
        annual_p = df.groupby("year")["precipitation"].sum().dropna()
        if len(annual_p) >= 5:
            x = annual_p.index.values.astype(float)
            y = annual_p.values
            slope, _, r, p, _ = stats.linregress(x, y)
            direction = "증가" if slope > 0 else "감소"
            sig = "통계적으로 유의한" if p < 0.05 else "뚜렷하지 않은"
            lines.append(
                f"[강수] {period} 연강수량은 10년당 {abs(slope * 10):.0f}mm "
                f"{direction} 추세이나 {sig} 변화임 (R²={r**2:.2f})."
            )
            lines.append(
                f"  분석 기간 평균 연강수량: {annual_p.mean():.0f}mm "
                f"(최소 {annual_p.min():.0f}mm [{int(annual_p.idxmin())}년], "
                f"최대 {annual_p.max():.0f}mm [{int(annual_p.idxmax())}년])"
            )

    return lines


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Matplotlib 차트 → BytesIO
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _chart_temp_trend(df: pd.DataFrame) -> bytes | None:
    """연평균 기온 추이 차트."""
    if "year" not in df.columns or not _has(df, "temp_avg"):
        return None

    annual = df.groupby("year")["temp_avg"].mean().dropna()
    if len(annual) < 3:
        return None

    fig, ax = plt.subplots(figsize=(CHART_WIDTH, CHART_HEIGHT_INCH), dpi=CHART_DPI)
    ax.plot(annual.index, annual.values, "o-", color="#2E75B6", lw=2, label="연평균기온")

    # 추세선
    x = annual.index.values.astype(float)
    slope, intercept, *_ = stats.linregress(x, annual.values)
    ax.plot(x, slope * x + intercept, "--", color="#E74C3C", lw=1.5, label=f"추세선 ({slope*10:+.2f}°C/10년)")

    ax.set_xlabel("연도", fontproperties=KR_FONT, fontsize=11)
    ax.set_ylabel("기온 (°C)", fontproperties=KR_FONT, fontsize=11)
    ax.set_title("연평균 기온 추이", fontproperties=KR_FONT, fontsize=13, fontweight="bold")
    ax.legend(prop={"family": KR_FONT}, fontsize=10)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=CHART_DPI)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _chart_precip_trend(df: pd.DataFrame) -> bytes | None:
    """연강수량 추이 차트."""
    if "year" not in df.columns or not _has(df, "precipitation"):
        return None

    annual = df.groupby("year")["precipitation"].sum().dropna()
    if len(annual) < 3:
        return None

    fig, ax = plt.subplots(figsize=(CHART_WIDTH, CHART_HEIGHT_INCH), dpi=CHART_DPI)
    ax.bar(annual.index, annual.values, color="#3498DB", alpha=0.7, label="연강수량")

    # 이동평균
    if len(annual) >= 5:
        ma5 = annual.rolling(5, min_periods=3).mean()
        ax.plot(annual.index, ma5.values, "-", color="#E74C3C", lw=2, label="5년 이동평균")

    ax.set_xlabel("연도", fontproperties=KR_FONT, fontsize=11)
    ax.set_ylabel("강수량 (mm)", fontproperties=KR_FONT, fontsize=11)
    ax.set_title("연강수량 추이", fontproperties=KR_FONT, fontsize=13, fontweight="bold")
    ax.legend(prop={"family": KR_FONT}, fontsize=10)
    ax.grid(True, alpha=0.3, axis="y")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=CHART_DPI)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _chart_mv_curves(df: pd.DataFrame) -> bytes | None:
    """MV 커브 (누적강수 비율) matplotlib 차트."""
    if "station_name" not in df.columns or not _has(df, "precipitation"):
        return None
    if "year" not in df.columns or "month" not in df.columns:
        return None

    stations = df["station_name"].dropna().unique()
    n_stations = len(stations)
    if n_stations == 0:
        return None

    cols_per_row = min(n_stations, 2)
    n_rows = (n_stations + cols_per_row - 1) // cols_per_row
    fig, axes = plt.subplots(
        n_rows, cols_per_row,
        figsize=(CHART_WIDTH * cols_per_row / 2 + 1, CHART_HEIGHT_INCH * n_rows),
        dpi=CHART_DPI,
        squeeze=False,
    )
    axes_flat = axes.flatten()

    month_labels = [f"{m}월" for m in range(1, 13)]

    for ax_idx, stn in enumerate(stations):
        ax = axes_flat[ax_idx]
        sdf = df[df["station_name"] == stn]
        monthly = sdf.groupby(["year", "month"])["precipitation"].sum().unstack("month")
        monthly = monthly.reindex(columns=range(1, 13))
        annual = monthly.sum(axis=1)
        ratio = monthly.cumsum(axis=1).div(annual, axis=0)

        years = sorted(ratio.index.tolist())
        recent_10 = [y for y in years if y >= years[-1] - 9]
        recent_30 = [y for y in years if y >= years[-1] - 29]

        # 개별 연도 (회색)
        for yr in years:
            row = ratio.loc[yr].values.tolist()
            ax.plot(range(1, 13), row, color="lightgray", lw=0.7, alpha=0.5)

        # 전체 평균
        mean_all = ratio.mean().values
        ax.plot(range(1, 13), mean_all, color="#1F4E79", lw=2.2,
                label=f"전체평균 ({years[0]}~{years[-1]})", marker="o", ms=4)

        # 10년 평균
        if len(recent_10) >= 3:
            m10 = ratio.loc[recent_10].mean().values
            ax.plot(range(1, 13), m10, color="#E74C3C", lw=2.0, ls="--",
                    label=f"최근10년 ({recent_10[0]}~{recent_10[-1]})", marker="^", ms=4)

        # 30년 평균
        if len(recent_30) >= 10 and len(recent_30) != len(years):
            m30 = ratio.loc[recent_30].mean().values
            ax.plot(range(1, 13), m30, color="#27AE60", lw=2.0, ls=":",
                    label=f"최근30년 ({recent_30[0]}~{recent_30[-1]})", marker="s", ms=4)

        ax.set_xticks(range(1, 13))
        ax.set_xticklabels(month_labels, fontproperties=KR_FONT, fontsize=8)
        ax.set_yticks([0, 0.2, 0.4, 0.6, 0.8, 1.0])
        ax.set_yticklabels(["0%", "20%", "40%", "60%", "80%", "100%"], fontsize=8)
        ax.set_ylim(-0.02, 1.05)
        ax.set_title(f"{stn} MV 커브", fontproperties=KR_FONT, fontsize=11, fontweight="bold")
        ax.set_xlabel("월", fontproperties=KR_FONT, fontsize=9)
        ax.set_ylabel("누적강수량 / 연강수량", fontproperties=KR_FONT, fontsize=9)
        ax.legend(prop={"family": KR_FONT}, fontsize=8, loc="upper left")
        ax.grid(True, alpha=0.3)

    # 남은 subplot 숨기기
    for ax_idx in range(n_stations, len(axes_flat)):
        axes_flat[ax_idx].set_visible(False)

    plt.suptitle("MV 커브 (누적강수 비율 곡선)", fontproperties=KR_FONT,
                 fontsize=13, fontweight="bold", y=1.01)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=CHART_DPI, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _chart_annual_temp_precip(df: pd.DataFrame) -> bytes | None:
    """그림 2-1 스타일: 연도별 기온 + 강수 (2개 subplot)."""
    if "year" not in df.columns:
        return None
    has_t = _has(df, "temp_avg") or _has(df, "temp_max") or _has(df, "temp_min")
    has_p = _has(df, "precipitation")
    if not has_t and not has_p:
        return None

    stations = df["station_name"].dropna().unique().tolist() if "station_name" in df.columns else []
    multi_station = len(stations) > 1
    colors = plt.cm.Set1.colors if multi_station else [("#2E75B6", "#E67E22", "#95A5A6")]

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(CHART_WIDTH, CHART_HEIGHT_INCH * 2.2), dpi=CHART_DPI)

    # ─ 상단: 기온 라인 차트 ─
    if has_t:
        if multi_station:
            for idx, stn in enumerate(stations):
                sdf = df[df["station_name"] == stn]
                c = colors[idx % len(colors)]
                if _has(sdf, "temp_avg"):
                    avg = sdf.groupby("year")["temp_avg"].mean().dropna()
                    ax1.plot(avg.index, avg.values, "o-", color=c, lw=2, ms=5, label=f"{stn} 평균기온(℃)")
        else:
            sdf = df
            t_color, tmax_color, tmin_color = "#2E75B6", "#E67E22", "#95A5A6"
            if _has(sdf, "temp_avg"):
                avg = sdf.groupby("year")["temp_avg"].mean().dropna()
                ax1.plot(avg.index, avg.values, "o-", color=t_color, lw=2, ms=5, label="평균기온(℃)")
            if _has(sdf, "temp_max"):
                tmax = sdf.groupby("year")["temp_max"].mean().dropna()
                ax1.plot(tmax.index, tmax.values, "o-", color=tmax_color, lw=2, ms=5, label="평균최고기온(℃)")
            if _has(sdf, "temp_min"):
                tmin = sdf.groupby("year")["temp_min"].mean().dropna()
                ax1.plot(tmin.index, tmin.values, "o-", color=tmin_color, lw=2, ms=5, label="평균최저기온(℃)")

    ax1.set_xlabel("연도", fontproperties=KR_FONT, fontsize=10)
    ax1.set_ylabel("기온 (℃)", fontproperties=KR_FONT, fontsize=10)
    ax1.legend(prop={"family": KR_FONT}, fontsize=9)
    ax1.grid(True, alpha=0.3)

    # ─ 하단: 강수 막대+선 이중축 ─
    if has_p:
        if multi_station:
            bar_width = 0.8 / max(len(stations), 1)
            for idx, stn in enumerate(stations):
                sdf = df[df["station_name"] == stn]
                if not _has(sdf, "precipitation"):
                    continue
                annual_p = sdf.groupby("year")["precipitation"].sum().dropna()
                c = colors[idx % len(colors)]
                offset = (idx - len(stations) / 2 + 0.5) * bar_width
                ax2.bar(annual_p.index + offset, annual_p.values, width=bar_width,
                        color=c, alpha=0.7, label=f"{stn} 연강수량")
        else:
            annual_p = df.groupby("year")["precipitation"].sum().dropna()
            ax2.bar(annual_p.index, annual_p.values, color="#3498DB", alpha=0.7, label="연강수량")
            ax2_twin = ax2.twinx()
            monthly_avg = annual_p / 12
            ax2_twin.plot(annual_p.index, monthly_avg.values, "o-", color="#E67E22", lw=2, ms=5,
                          label="월평균강수량")
            ax2_twin.set_ylabel("월평균강수량 (mm)", fontproperties=KR_FONT, fontsize=10)
            lines2, labels2 = ax2.get_legend_handles_labels()
            lines3, labels3 = ax2_twin.get_legend_handles_labels()
            ax2.legend(lines2 + lines3, labels2 + labels3, prop={"family": KR_FONT}, fontsize=9)

    ax2.set_xlabel("연도", fontproperties=KR_FONT, fontsize=10)
    ax2.set_ylabel("연강수량 (mm)", fontproperties=KR_FONT, fontsize=10)
    if multi_station:
        ax2.legend(prop={"family": KR_FONT}, fontsize=9)
    ax2.grid(True, alpha=0.3)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=CHART_DPI, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _chart_monthly_boxplot(df: pd.DataFrame) -> bytes | None:
    """그림 2-2 스타일: 월별 박스플롯 (강수량 + 기온 2개 subplot)."""
    if "month" not in df.columns or "year" not in df.columns:
        return None
    has_p = _has(df, "precipitation")
    has_t = _has(df, "temp_avg")
    if not has_p and not has_t:
        return None

    # 복수 관측소면 첫 번째 관측소만 사용
    if "station_name" in df.columns and df["station_name"].nunique() > 1:
        first_stn = df["station_name"].dropna().unique()[0]
        df = df[df["station_name"] == first_stn]

    month_labels = ["1월", "2월", "3월", "4월", "5월", "6월",
                    "7월", "8월", "9월", "10월", "11월", "12월"]

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(CHART_WIDTH, CHART_HEIGHT_INCH * 2.2), dpi=CHART_DPI)

    # ─ 상단: 월별 강수량 박스플롯 ─
    if has_p:
        monthly_precip = df.groupby(["year", "month"])["precipitation"].sum()
        data_p = [
            monthly_precip[monthly_precip.index.get_level_values("month") == m].values
            for m in range(1, 13)
        ]
        bp1 = ax1.boxplot(data_p, patch_artist=True, flierprops={"marker": "o", "ms": 3})
        for patch in bp1["boxes"]:
            patch.set_facecolor("#3498DB")
        ax1.set_xticklabels(month_labels, fontproperties=KR_FONT, fontsize=9)
        ax1.set_ylabel("강수량 [mm]", fontproperties=KR_FONT, fontsize=10)
        ax1.set_title("월별 강수량 분포", fontproperties=KR_FONT, fontsize=11, fontweight="bold")
        ax1.grid(True, alpha=0.3)
    else:
        ax1.set_visible(False)

    # ─ 하단: 월별 기온 박스플롯 ─
    if has_t:
        monthly_temp = df.groupby(["year", "month"])["temp_avg"].mean()
        data_t = [
            monthly_temp[monthly_temp.index.get_level_values("month") == m].values
            for m in range(1, 13)
        ]
        bp2 = ax2.boxplot(data_t, patch_artist=True, flierprops={"marker": "o", "ms": 3})
        for patch in bp2["boxes"]:
            patch.set_facecolor("#E67E22")
        ax2.set_xticklabels(month_labels, fontproperties=KR_FONT, fontsize=9)
        ax2.set_ylabel("기온 [℃]", fontproperties=KR_FONT, fontsize=10)
        ax2.set_title("월별 기온 분포", fontproperties=KR_FONT, fontsize=11, fontweight="bold")
        ax2.grid(True, alpha=0.3)
    else:
        ax2.set_visible(False)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=CHART_DPI, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _chart_monthly_bar(df: pd.DataFrame) -> bytes | None:
    """월별 평균 기온·강수량 이중 축 차트."""
    if "month" not in df.columns:
        return None
    has_t = _has(df, "temp_avg")
    has_p = _has(df, "precipitation")
    if not has_t and not has_p:
        return None

    months = list(range(1, 13))
    month_labels = [f"{m}월" for m in months]

    fig, ax1 = plt.subplots(figsize=(CHART_WIDTH, CHART_HEIGHT_INCH), dpi=CHART_DPI)

    if has_p:
        monthly_p = df.groupby("month")["precipitation"].sum() / df["year"].nunique() if "year" in df.columns \
            else df.groupby("month")["precipitation"].sum()
        p_vals = [monthly_p.get(m, 0) for m in months]
        ax1.bar(months, p_vals, color="#85C1E9", alpha=0.8, label="강수량(mm)")
        ax1.set_ylabel("강수량 (mm)", fontproperties=KR_FONT, fontsize=10, color="#2980B9")
        ax1.tick_params(axis="y", labelcolor="#2980B9")

    if has_t:
        ax2 = ax1.twinx()
        monthly_t = df.groupby("month")["temp_avg"].mean()
        t_vals = [monthly_t.get(m, np.nan) for m in months]
        ax2.plot(months, t_vals, "o-", color="#E74C3C", lw=2, label="평균기온(°C)")
        ax2.set_ylabel("기온 (°C)", fontproperties=KR_FONT, fontsize=10, color="#E74C3C")
        ax2.tick_params(axis="y", labelcolor="#E74C3C")

    ax1.set_xticks(months)
    ax1.set_xticklabels(month_labels, fontproperties=KR_FONT, fontsize=9)
    ax1.set_title("월별 기온·강수량", fontproperties=KR_FONT, fontsize=13, fontweight="bold")

    lines1, labels1 = ax1.get_legend_handles_labels()
    if has_t:
        lines2, labels2 = ax2.get_legend_handles_labels()
        lines1 += lines2
        labels1 += labels2
    ax1.legend(lines1, labels1, prop={"family": KR_FONT}, fontsize=9, loc="upper left")
    ax1.grid(True, alpha=0.3, axis="y")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=CHART_DPI)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# DOCX 스타일 헬퍼
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _set_cell_bg(cell, hex_color: str) -> None:
    """표 셀 배경색 설정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _style_header_row(table, hex_color: str = HEADER_BG) -> None:
    """표 첫 행을 헤더 스타일로 처리."""
    for cell in table.rows[0].cells:
        _set_cell_bg(cell, hex_color)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)


def _add_heading(doc: "Document", text: str, level: int = 1) -> None:
    """헤딩 추가."""
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = TITLE_COLOR if level == 1 else ACCENT_COLOR


def _add_df_table(doc: "Document", df: pd.DataFrame, header_bg: str = HEADER_BG) -> None:
    """DataFrame을 DOCX 표로 삽입."""
    if df.empty:
        doc.add_paragraph("(데이터 없음)")
        return

    cols = list(df.columns)
    table = doc.add_table(rows=1 + len(df), cols=len(cols))
    table.style = "Table Grid"

    # 헤더
    for j, col in enumerate(cols):
        cell = table.cell(0, j)
        cell.text = str(col)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    _style_header_row(table, header_bg)

    # 데이터 행
    for i, (_, row) in enumerate(df.iterrows()):
        bg = ROW_BG_EVEN if i % 2 == 0 else "FFFFFF"
        for j, col in enumerate(cols):
            val = row[col]
            cell = table.cell(i + 1, j)
            if isinstance(val, float):
                cell.text = f"{val:.2f}"
            else:
                cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if bg != "FFFFFF":
                _set_cell_bg(cell, bg)


def _add_chart_image(doc: "Document", img_bytes: bytes | None, caption: str = "") -> None:
    """차트 이미지 삽입."""
    if img_bytes is None:
        return
    buf = io.BytesIO(img_bytes)
    doc.add_picture(buf, width=Inches(6.0))
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 메인 DOCX 생성 클래스
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class DocxReportGenerator:
    """DOCX 보고서 생성 (python-docx 기반)."""

    def generate_docx(
        self,
        df: pd.DataFrame,
        monthly_df: pd.DataFrame | None = None,
        climate_df: pd.DataFrame | None = None,
    ) -> bytes:
        """
        DOCX 보고서 생성 후 bytes 반환.
        df: 일별 필터링된 DataFrame (WeatherDataProcessor 결과)
        """
        if not _DOCX_AVAILABLE:
            raise ImportError("python-docx 패키지가 설치되지 않았습니다. requirements.txt에 python-docx 추가 필요.")

        doc = Document()

        # 기본 여백 설정
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.5)

        self._write_cover(doc, df)

        # 기상개황 (보고서 예시 스타일)
        self._write_weather_overview(doc, df)
        self._write_chart_annual(doc, df)
        doc.add_page_break()
        self._write_climate_monthly_table(doc, df)
        self._write_chart_monthly(doc, df)
        doc.add_page_break()

        # 기존 섹션
        self._write_summary(doc, df)
        self._write_charts(doc, df)
        self._write_monthly_table(doc, df)
        self._write_annual_table(doc, df)
        self._write_interpretation(doc, df)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def generate_filename(self) -> str:
        today = datetime.now().strftime("%Y%m%d")
        return f"기상분석보고서_{today}.docx"

    # ── 섹션별 작성 메서드 ────────────────────────────────

    def _write_cover(self, doc: "Document", df: pd.DataFrame) -> None:
        """1. 표지."""
        # 제목
        title = doc.add_heading("", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("기상자료 분석 보고서")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = TITLE_COLOR

        doc.add_paragraph()  # 여백

        # 부제목
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run("Weather Data Analysis Report")
        r.font.size = Pt(14)
        r.font.color.rgb = ACCENT_COLOR
        r.font.italic = True

        doc.add_paragraph()
        doc.add_paragraph()

        # 메타 정보 테이블
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = "Table Grid"

        meta_rows = [
            ("분석 기간", self._period_str(df)),
            ("관측소", self._station_str(df)),
            ("총 관측일수", f"{len(df):,}일"),
            ("보고서 생성일", datetime.now().strftime("%Y년 %m월 %d일")),
        ]

        for i, (label, val) in enumerate(meta_rows):
            cell_l = info_table.cell(i, 0)
            cell_r = info_table.cell(i, 1)
            cell_l.text = label
            cell_r.text = val
            _set_cell_bg(cell_l, "1F4E79")
            cell_l.paragraphs[0].runs[0].font.bold = True
            cell_l.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell_l.paragraphs[0].runs[0].font.size = Pt(11)
            cell_r.paragraphs[0].runs[0].font.size = Pt(11)

        doc.add_page_break()

    def _write_weather_overview(self, doc: "Document", df: pd.DataFrame) -> None:
        """기상개황 — 표 2-1 스타일."""
        _add_heading(doc, "1. 기상개황", 1)
        _add_heading(doc, "1.1 연도별 기상개황", 2)
        doc.add_paragraph("분석 기간 중 연도별 기상요소 현황입니다.")

        if "year" not in df.columns:
            doc.add_paragraph("(연도 데이터 없음)")
            return

        stations = (
            df["station_name"].dropna().unique().tolist()
            if "station_name" in df.columns
            else ["전체"]
        )
        years = sorted(df["year"].dropna().unique().astype(int).tolist())
        year_cols = [str(y) for y in years]

        # 연도 범위 각주용
        start_yr = years[0] if years else ""
        end_yr = years[-1] if years else ""

        for stn in stations:
            _add_heading(doc, f"{stn} 관측소 기상개황", 3)

            sdf = df[df["station_name"] == stn] if "station_name" in df.columns else df

            rows = []

            # 평균기온 3행
            for sub, col, agg_fn in [
                ("평균", "temp_avg", "mean"),
                ("최고", "temp_max", "mean"),
                ("최저", "temp_min", "mean"),
            ]:
                if _has(sdf, col):
                    annual = sdf.groupby("year")[col].agg(agg_fn).round(1)
                    label = "평균기온(℃)" if sub == "평균" else ""
                    vals = {str(y): f"{annual.get(y, float('nan')):.1f}" if y in annual.index else "-"
                            for y in years}
                    mean_val = annual.mean()
                    rows.append({"항목": label, "구분": sub,
                                 **vals, "평균": f"{mean_val:.1f}"})

            # 평균월강수량
            if _has(sdf, "precipitation"):
                annual_p = sdf.groupby("year")["precipitation"].sum().round(1)
                monthly_p = (annual_p / 12).round(1)
                vals = {str(y): f"{monthly_p.get(y, float('nan')):.1f}" if y in monthly_p.index else "-"
                        for y in years}
                rows.append({"항목": "평균월강수량(mm)", "구분": "",
                             **vals, "평균": f"{monthly_p.mean():.1f}"})

            # 평균풍속
            if _has(sdf, "wind_speed"):
                annual_ws = sdf.groupby("year")["wind_speed"].mean().round(1)
                vals = {str(y): f"{annual_ws.get(y, float('nan')):.1f}" if y in annual_ws.index else "-"
                        for y in years}
                rows.append({"항목": "평균풍속(m/s)", "구분": "",
                             **vals, "평균": f"{annual_ws.mean():.1f}"})

            # 최대순간풍속
            if _has(sdf, "wind_max"):
                annual_wm = sdf.groupby("year")["wind_max"].max().round(1)
                vals = {str(y): f"{annual_wm.get(y, float('nan')):.1f}" if y in annual_wm.index else "-"
                        for y in years}
                rows.append({"항목": "최대순간풍속(m/s)", "구분": "",
                             **vals, "평균": f"{annual_wm.mean():.1f}"})

            if rows:
                overview_df = pd.DataFrame(rows, columns=["항목", "구분"] + year_cols + ["평균"])
                _add_df_table(doc, overview_df)
                doc.add_paragraph()

            # 월별 강수 피벗표
            if _has(sdf, "precipitation") and "month" in sdf.columns:
                _add_heading(doc, f"{stn} 관측소 월별 강수량 현황", 3)
                mp = sdf.groupby(["year", "month"])["precipitation"].sum().unstack("month")
                mp = mp.reindex(columns=range(1, 13))
                mp.columns = [f"{m}월" for m in range(1, 13)]
                mp["합계"] = mp.sum(axis=1)
                mp.index.name = "연도"
                mp = mp.reset_index()
                mp["연도"] = mp["연도"].astype(int).astype(str)

                # 하단 통계행
                precip_cols = [f"{m}월" for m in range(1, 13)] + ["합계"]
                numeric_mp = mp.set_index("연도")[precip_cols]
                avg_row = {"연도": "평균"}
                max_row = {"연도": "최대"}
                min_row = {"연도": "최소"}
                for c in precip_cols:
                    avg_row[c] = f"{numeric_mp[c].mean():.1f}"
                    max_row[c] = f"{numeric_mp[c].max():.1f}"
                    min_row[c] = f"{numeric_mp[c].min():.1f}"

                mp_str = mp.copy()
                for c in precip_cols:
                    mp_str[c] = mp_str[c].apply(
                        lambda v: f"{v:.1f}" if pd.notna(v) else "-"
                    )
                stat_rows = pd.DataFrame([avg_row, max_row, min_row])
                mp_final = pd.concat([mp_str, stat_rows], ignore_index=True)
                _add_df_table(doc, mp_final)
                doc.add_paragraph()

        note = doc.add_paragraph(f"※ 자료 : 기상자료개방포털({start_yr}~{end_yr})")
        note.runs[0].font.italic = True
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    def _write_chart_annual(self, doc: "Document", df: pd.DataFrame) -> None:
        """연도별 기상 현황 차트 — 그림 2-1."""
        img = _chart_annual_temp_precip(df)
        if img:
            _add_heading(doc, "1.2 연도별 기상 현황", 2)
            doc.add_paragraph("연도별 기온 추이(상) 및 강수량 현황(하)입니다.")
            _add_chart_image(doc, img, "그림 1. 연평균 기상 개황도")
            doc.add_paragraph()

    def _write_climate_monthly_table(self, doc: "Document", df: pd.DataFrame) -> None:
        """월별 기후통계표 — 표 2-2 스타일."""
        _add_heading(doc, "2. 월별 기후통계", 1)
        _add_heading(doc, "2.1 월별 기후통계표", 2)

        if "month" not in df.columns:
            doc.add_paragraph("(월 데이터 없음)")
            return

        if "year" not in df.columns:
            doc.add_paragraph("(연도 데이터 없음)")
            return

        stations = (
            df["station_name"].dropna().unique().tolist()
            if "station_name" in df.columns
            else ["전체"]
        )
        month_labels = [f"{m}월" for m in range(1, 13)]
        start_yr = int(df["year"].min()) if "year" in df.columns else ""
        end_yr = int(df["year"].max()) if "year" in df.columns else ""

        for stn in stations:
            _add_heading(doc, f"{stn} 관측소 월별 기후통계", 3)
            sdf = df[df["station_name"] == stn] if "station_name" in df.columns else df

            def _monthly_stat(col, agg_fn):
                if not _has(sdf, col):
                    return None
                grp = sdf.groupby("month")[col].agg(agg_fn).reindex(range(1, 13))
                return grp

            def _monthly_extreme(col, agg_fn):
                if not _has(sdf, col):
                    return None
                return sdf.groupby("month")[col].agg(agg_fn).reindex(range(1, 13))

            def _fmt(series, last_label):
                if series is None:
                    return ["-"] * 12 + ["-"]
                vals = [f"{v:.1f}" if pd.notna(v) else "-" for v in series.values]
                summary = series.mean() if last_label == "평균" else series.sum()
                return vals + [f"{summary:.1f}" if pd.notna(summary) else "-"]

            rows = []

            col_header = ["구분", "소구분"] + month_labels + ["평균/합계"]

            # 기온
            temp_defs = [
                ("기온(℃)", "평균", "temp_avg", "mean", "평균"),
                ("", "평균최고", "temp_max", "mean", "평균"),
                ("", "최고극값", "temp_max", "max", "평균"),
                ("", "평균최저", "temp_min", "mean", "평균"),
                ("", "최저극값", "temp_min", "min", "평균"),
            ]
            for cat, sub, col, fn, last_lbl in temp_defs:
                series = _monthly_extreme(col, fn) if fn in ("max", "min") else _monthly_stat(col, fn)
                rows.append([cat, sub] + _fmt(series, last_lbl))

            # 강수량
            precip_defs = [
                ("강수량(mm)", "평균", "precipitation", lambda s: s.sum() / sdf["year"].nunique()
                    if "year" in sdf.columns else s.sum(), "합계"),
                ("", "최대", "precipitation", "max", "합계"),
                ("", "최소", "precipitation", "min", "합계"),
            ]
            for cat, sub, col, fn, last_lbl in precip_defs:
                if not _has(sdf, col):
                    rows.append([cat, sub] + ["-"] * 13)
                    continue
                if callable(fn):
                    series = sdf.groupby("month")[col].agg(fn).reindex(range(1, 13))
                else:
                    series = sdf.groupby("month")[col].agg(fn).reindex(range(1, 13))
                rows.append([cat, sub] + _fmt(series, last_lbl))

            # 상대습도
            humidity_defs = [
                ("상대습도(%)", "평균", "humidity", "mean", "평균"),
                ("", "최소", "humidity", "min", "평균"),
            ]
            for cat, sub, col, fn, last_lbl in humidity_defs:
                series = _monthly_stat(col, fn)
                rows.append([cat, sub] + _fmt(series, last_lbl))

            # 일조시간
            sunshine_defs = [
                ("일조시간(hr)", "", "sunshine_hours", "sum", "합계"),
            ]
            for cat, sub, col, fn, last_lbl in sunshine_defs:
                series = _monthly_stat(col, fn) if _has(sdf, col) else None
                rows.append([cat, sub] + _fmt(series, last_lbl))

            # 바람
            wind_defs = [
                ("바람(m/s)", "평균풍속", "wind_speed", "mean", "평균"),
                ("", "최대풍속", "wind_max", "max", "평균"),
            ]
            for cat, sub, col, fn, last_lbl in wind_defs:
                series = _monthly_stat(col, fn)
                rows.append([cat, sub] + _fmt(series, last_lbl))

            climate_df = pd.DataFrame(rows, columns=col_header)
            _add_df_table(doc, climate_df)
            doc.add_paragraph()

        note = doc.add_paragraph(f"※ 자료 : 기상자료개방포털({start_yr}~{end_yr})")
        note.runs[0].font.italic = True
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    def _write_chart_monthly(self, doc: "Document", df: pd.DataFrame) -> None:
        """월별 기상 현황 박스플롯 차트 — 그림 2-2."""
        img = _chart_monthly_boxplot(df)
        if img:
            _add_heading(doc, "2.2 월별 기상 현황", 2)
            doc.add_paragraph("월별 강수량(상) 및 기온(하) 분포입니다.")
            _add_chart_image(doc, img, "그림 2. 월 평균 기상 현황도")
            doc.add_paragraph()

    def _write_summary(self, doc: "Document", df: pd.DataFrame) -> None:
        """2. 주요 지표 요약."""
        _add_heading(doc, "1. 주요 기상 지표 요약", 1)

        summary = _compute_summary(df)
        if not summary:
            doc.add_paragraph("요약할 데이터가 없습니다.")
            return

        table = doc.add_table(rows=1 + len(summary), cols=3)
        table.style = "Table Grid"

        # 헤더
        headers = ["지표", "값", "단위"]
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = h
        _style_header_row(table)

        for i, (label, (val, unit)) in enumerate(summary.items()):
            bg = ROW_BG_EVEN if i % 2 == 0 else "FFFFFF"
            table.cell(i + 1, 0).text = label
            table.cell(i + 1, 1).text = f"{val:.1f}" if isinstance(val, float) else str(val)
            table.cell(i + 1, 2).text = unit
            for j in range(3):
                table.cell(i + 1, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if bg != "FFFFFF":
                    _set_cell_bg(table.cell(i + 1, j), bg)

        doc.add_paragraph()

    def _write_charts(self, doc: "Document", df: pd.DataFrame) -> None:
        """3. 핵심 차트."""
        _add_heading(doc, "2. 기상 추이 분석", 1)

        img1 = _chart_temp_trend(df)
        if img1:
            _add_heading(doc, "2.1 연평균 기온 추이", 2)
            _add_chart_image(doc, img1, "그림 1. 연평균 기온 추이 및 선형 추세선")
            doc.add_paragraph()

        img2 = _chart_precip_trend(df)
        if img2:
            _add_heading(doc, "2.2 연강수량 추이", 2)
            _add_chart_image(doc, img2, "그림 2. 연강수량 추이 및 5년 이동평균")
            doc.add_paragraph()

        img3 = _chart_monthly_bar(df)
        if img3:
            _add_heading(doc, "2.3 월별 기온·강수량", 2)
            _add_chart_image(doc, img3, "그림 3. 월별 평균 기온 및 강수량 분포")
            doc.add_paragraph()

        img4 = _chart_mv_curves(df)
        if img4:
            _add_heading(doc, "2.4 MV 커브 (누적강수 비율 곡선)", 2)
            doc.add_paragraph(
                "MV 커브는 월별 누적강수량을 연강수량으로 나눈 비율(0~1)을 나타냅니다. "
                "곡선의 기울기가 급한 구간이 강수 집중 시기이며, "
                "여름철(6~9월)에 집중되는 한국의 계절 강수 패턴을 시각적으로 확인할 수 있습니다."
            )
            _add_chart_image(doc, img4, "그림 4. MV 커브 — 회색: 개별 연도, 색선: 기간별 평균")
            doc.add_paragraph()

        doc.add_page_break()

    def _write_monthly_table(self, doc: "Document", df: pd.DataFrame) -> None:
        """4. 월별 통계표."""
        _add_heading(doc, "3. 월별 기상 통계", 1)
        monthly = _compute_monthly_stats(df)
        if monthly.empty:
            doc.add_paragraph("월별 통계를 계산할 데이터가 없습니다.")
            return
        _add_df_table(doc, monthly)
        doc.add_paragraph()

    def _write_annual_table(self, doc: "Document", df: pd.DataFrame) -> None:
        """5. 연도별 통계표."""
        _add_heading(doc, "4. 연도별 기상 통계", 1)
        annual = _compute_annual_stats(df)
        if annual.empty:
            doc.add_paragraph("연도별 통계를 계산할 데이터가 없습니다.")
            return
        _add_df_table(doc, annual)
        doc.add_paragraph()

    def _write_interpretation(self, doc: "Document", df: pd.DataFrame) -> None:
        """6. 자동 해석 텍스트."""
        _add_heading(doc, "5. 분석 결과 해석", 1)

        lines = _auto_interpret(df)
        if not lines:
            doc.add_paragraph("데이터가 부족하여 자동 해석을 생성할 수 없습니다.")
            return

        for line in lines:
            p = doc.add_paragraph()
            if line.startswith("["):
                # 소제목 (예: [기온], [강수])
                run = p.add_run(line)
                run.font.bold = True if line.startswith("[") else False
                run.font.size = Pt(10)
            else:
                run = p.add_run(line)
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.25)

        # 면책 문구
        doc.add_paragraph()
        note = doc.add_paragraph(
            "※ 본 보고서는 기상청 ASOS/AWS 일자료를 기반으로 자동 생성되었습니다. "
            "공식 기후통계는 기상청 기후통계분석 서비스(https://data.kma.go.kr)를 참조하시기 바랍니다."
        )
        note.runs[0].font.italic = True
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # ── 유틸 ─────────────────────────────────────────────

    def _period_str(self, df: pd.DataFrame) -> str:
        if "date" not in df.columns:
            return "알 수 없음"
        d_min = df["date"].min()
        d_max = df["date"].max()
        return f"{d_min.strftime('%Y년 %m월 %d일')} ~ {d_max.strftime('%Y년 %m월 %d일')}"

    def _station_str(self, df: pd.DataFrame) -> str:
        if "station_name" not in df.columns:
            return "알 수 없음"
        stations = sorted(df["station_name"].dropna().unique())
        return ", ".join(stations)
